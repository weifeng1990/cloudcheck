from docx.shared import Mm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 创建表格,默认行距为1cm
def createTable(document, row, col):
    # table = document.add_table(row, col, style='Medium Grid 1 Accent 1')
    table = document.add_table(row, col, style='Table Grid')
    table.style.font.name = u'宋体'
    table.style.font.size = Pt(11)
    for i in table.rows[0].cells:
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="B0C4DE"/>'.format(nsdecls('w')))
        i._tc.get_or_add_tcPr().append(shading_elm_2)
        del shading_elm_2
    # table = document.add_table(row, col, style='Medium Shading 2 Accent 1')
    for i in table.rows:
        i.height = Mm(10)
    return table

#
def osBasicDocument(document, list1):
    h1 = document.add_heading('Cloudos平台巡检结果')
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h2 = document.add_heading('1.Cloduos平台信息')
    t1 = createTable(document, 5, 2)
    t1.style.font.name = u'楷体'
    t1.cell(0, 0).text = "服务器型号"
    t1.cell(1, 0).text = "服务器规格"
    t1.cell(2, 0).text = "部署方式"
    t1.cell(3, 0).text = "集群节点数"
    t1.cell(4, 0).text = "版本号"
    for i in range(5):
        t1.cell(i, 1).text = list1[i]
    return

def osPlatDocument(document, list1, list2):
    h1 = document.add_heading("2.云管理平台状态及功能检查云管理平台状态及功能检查")
    count = 0
    text = str()
    for i in list2:
        if i:
            count += 1
    p1 = document.add_paragraph()
    run1 = p1.add_run("巡检小结：")
    run1.font.name = u'宋体'
    run1.font.size = Pt(11)
    text = "对CloudOS云管理平台进行巡检，巡检异常项数：" + (str)(count) + "；" + "正常项数：" + (str)(len(list2) - count)
    p2 = document.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.3)
    run2 = p2.add_run(text)
    run2.font.name = u'宋体'
    run2.font.size = Pt(11)
    t1 = createTable(document, 13, 4)
    t1.cell(0, 0).text = "检查内容"
    t1.cell(0, 1).text = "检查方法"
    t1.cell(0, 2).text = "检查结果"
    t1.cell(0, 3).text = "说明"
    t1.cell(1, 0).text = "服务器本地磁盘分区检查"
    t1.cell(1, 1).text = "登录各节点操作系统，执行命令检查分区是否正确"
    t1.cell(2, 0).text = "服务器可用空间检查"
    t1.cell(2, 1).text = "登录各节点操作系统，执行命令检查服务器本地磁盘及存储卷的利用率是否高于80%"
    t1.cell(3, 0).text = "服务器本地时间检查"
    t1.cell(3, 1).text = "登录Cluster节点和独立计算节点操作系统，执行命令查看各节点服务器与Master节点的时间是否同步"
    t1.cell(4, 0).text = "共享存储卷通断检查"
    t1.cell(4, 1).text = "登录各节点操作系统，执行命令检查是否有与共享存储卷相关的错误日志"
    t1.cell(5, 0).text = "容器镜像完整性检查"
    t1.cell(5, 1).text = "登录各节点操作系统，使用命令检查容器镜像完整"
    t1.cell(6, 0).text = "cloudos各节点的cpu利用率检查"
    t1.cell(6, 1).text = "ssh登录各节点使用top查看cpu利用率是否查过80%"
    t1.cell(7, 0).text = "cloudos各节点的内存利用率是否正常"
    t1.cell(7, 1).text = "ssh登录各节点使用free查看内存利用率是否超过80%"
    t1.cell(8, 0).text = "容器状态"
    t1.cell(8, 1).text = "查看容器状态是否正常"
    t1.cell(9, 0).text = "容器分布检查"
    t1.cell(9, 1).text = "登录各Master节点操作系统，使用命令检查所有容器是否均匀的运行在集群的各节点上"
    t1.cell(10, 0).text = "关键服务状态检查"
    t1.cell(10, 1).text = "登录各节点操作系统并进入相关容器内部，使用命令检查关键服务的状态是否正常（active (running)）"
    t1.cell(11, 0).text = "云主机状态检查"
    t1.cell(11, 1).text = "使用云管理员账户登录H3Cloud云管理平台，单击[计算与存储/主机]菜单项，在页面中查看是否有异常状态的云主机。"
    t1.cell(12, 0).text = "云硬盘状态检查"
    t1.cell(12, 1).text = "使用云管理员账户登录H3Cloud云管理平台，单击[计算与存储/硬盘]菜单项，在页面中查看是否有异常状态的云硬盘。"
    for i in range(12):
        if not list2[i]:
            t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
        else:
            run = t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
            run.font.color.rgb = RGBColor(255, 0, 0)
            t1.cell(i + 1, 3).paragraphs[0].add_run(list2[i])
    return

def osBasicCheck(document, osInfo):
    list1 = ['' for n in range(5)]
    list1[0] = osInfo['productVersion']
    list1[1] = osInfo['deviceDmide']
    if len(osInfo['nodeInfo']) > 1:
        list1[2] = '集群'
    else:
        list1[2] = '单机'
    list1[3] = (str)(len(osInfo['nodeInfo']))
    list1[4] = osInfo['version']
    osBasicDocument(document, list1)
    return


def osPlatCheck(document, osInfo):
    list1 = list()
    list2 = ['' for n in range(12)]
    for i in osInfo['nodeInfo']:
        #检查分区是否合规
        temp = str()
        for j in  i['diskCapacity']:
            if j['name'] == 'centos-root':
                if j['capacity'] < 201:
                    if not temp:
                        temp = "\n主机节点" + i['hostName'] + "如下分区空间不合规：" + "centos-root"
                    else:
                        temp += '、centos-root'
            elif j['name'] == 'centos-swap':
                if j['capacity'] < 33.8:
                    if not temp:
                        temp = "\n主机节点" + i['hostName'] + "如下分区空间不合规：" + "centos-swap"
                    else:
                        temp += '、centos-swap'
            elif j['name'] == 'centos-metadata':
                if j['capacity'] < 5.3:
                    if not temp:
                        temp = "\n主机节点" + i['hostName'] + "如下分区空间不合规：" + "centos-metadata"
                    else:
                        temp += '、centos-metadata'
            elif j['name'] == 'centos-data':
                if j['capacity'] < 296:
                    if not temp:
                        temp = "\n主机节点" + i['hostName'] + "如下分区空间不合规：" + "centos-data"
                    else:
                        temp += '、centos-data'
        list2[0] += temp
        del temp

        #磁盘利用率
        temp1 = str()
        for k in i['diskRate']:
            if k['rate'] > 0.8:
                if not temp1:
                    temp1 = "\n主机节点" + i['hostName'] + "如下磁盘利用率超过过80%：" + k['name']
                else:
                    temp1 += "、" + k['name']
        list2[1] += temp1
        del temp1

        #各节点ntp时间是否与主节点偏移过大
        if i['ntpOffset'] > 10:
            if not list2[2]:
                list2[2] = "ntp时间与主节点不同步的主机如下：" + i['hostName']
            else:
                list2[2] += "、" + i['hostName']

        #共享存储是否正常
        if not i['shareStorError']:
            if not list2[3]:
                list2[3] = "共享存储异常的节点如下：" + i['hostName']
            else:
                list2[3] = "、" + i['hostName']

        #容器镜像完整性
        if i['hostName'] == osInfo['masterIP']:
            if len(i['images']) != 0:
                list2[4] += "\n主节点" + i['hostName'] + "缺少如下镜像："
                for k in i['images']:
                    list2[4] += "\t" + k
        else:
            if len(i['images']) != 0:
                list2[4] += "\n节点" + i['hostName'] + "缺少如下镜像："
                for k in i['images']:
                    list2[4] += "\t" + k

        #节点cpu利用率
        if i['cpuRate'] > 0.8:
            if not list2[5]:
                list2[5] = "\ncpu利用率大于80%节点如下:" + i['hostName']
            else:
                list2[5] += "、" + i['hostName']

        #节点内存利用率
        if i['memRate'] > 0.8:
            if not list2[6]:
                list2[6] = "\n内存利用率大于80%节点如下:" + i['hostName']
            else:
                list2[6] += "、" + i['hostName']

    #容器状态
    for i in osInfo['ctainrState']:
        if i['status'] != 'Running':
            if not list2[7]:
                list2[7] = '状态异常容器pod如下：' + i['name']
            else:
                list2[7] += '、' + i['name']

    # k8s集群容器分布是否均匀
    if not osInfo['ctainrLB']:
        list2[8] = "k8s集群容器分布不均匀"

    # 容器关键服务检查
    str1 = ''
    for j in osInfo['serviceStatus'].keys():
        str2 = ''
        for i in osInfo['serviceStatus'][j]:
            if not i['status']:
                print("#########", j, i['name'], i['status'])
                if not str2:
                    str2 = "\nPOD " + j +'如下服务异常：' + i['name']
                else:
                    str2 += "、" + i['name']
            else:
                continue
        if not str2:
            str1 += (str2 + ";")
    list2[9] = str1

    # 云主机
    for i in osInfo['vmStatus']:
        if i['status'] != "ACTIVE" and i['status'] != "SHUTOFF":
            if not list2[10]:
                list2[10] = '状态异常云主机如下：' + i['name']
            else:
                list2[10] += '、' + i['name']

    # 云硬盘
    for i in osInfo['vDiskStatus']:
        if i['status'] != 'available' and i['status'] != 'in-use':
            if not list2[11]:
                list2[11] = '状态异常云硬盘如下：' + i['name']
            else:
                list2[11] += '、' + i['name']

    for i in list2:
        if not i:
            list1.append("正常")
        else:
            list1.append("异常")

    osPlatDocument(document,list1,list2)
    del list1, list2
    return




