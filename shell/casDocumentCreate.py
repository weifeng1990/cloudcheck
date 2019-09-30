from docx.shared import Mm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

##
#cloudos2.0 api端口为9000
#cloudos3.0 api端口为8000

# 创建表格,默认行距为1cm
def createTable(document, row, col):
    # table = document.add_table(row, col, style='Medium Grid 1 Accent 1')
    table = document.add_table(row, col, style='Table Grid')
    table.style.font.name = u'宋体'
    table.style.font.size = Pt(11)
    for i in table.rows[0].cells:
        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="B0C4DE"/>'.format(nsdecls('w')))
        i._tc.get_or_add_tcPr().append(shading_elm_2)
        del shading_elm_2
    # table = document.add_table(row, col, style='Medium Shading 2 Accent 1')
    for i in table.rows:
        i.height = Mm(10)
    return table


# 合并单元格,返回单元格地址
def mergeCell(table, beginRow, beginCol, endRow, endCol):
    c1 = table.cell(beginRow, beginCol)
    c2 = table.cell(endRow, endCol)
    return c1.merge(c2)


# document = openDocument(r'cas.docx')
# serverList为参数列表，包含6个参数，从参数0-6分别为：服务器型号、服务器规格、CAS版本、CVM部署方式、
# S1020V版本、是否使用临时license
def casBasicDocument(document, list1):
    h1 = document.add_heading('CAS平台巡检结果')
    h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h2 = document.add_heading('1.CAS平台基本信息')
    t1 = createTable(document, 7, 2)

    # 初始化表格
    t1.cell(0, 0).text = "巡检项"
    t1.cell(0, 1).text = "参数"
    t1.cell(1, 0).text = "服务器型号"
    t1.cell(2, 0).text = "服务格器规"
    t1.cell(3, 0).text = "CAS版本号"
    t1.cell(4, 0).text = "CVM部署方式"
    t1.cell(5, 0).text = "S1020V版本号"
    t1.cell(6, 0).text = "是否有使用临时license"
    # 参数赋值
    for i in range(6):
        run = t1.cell(i + 1, 1).paragraphs[0].add_run(list1[i])
        # run.font.name = '宋体'
        # run.font.size = Pt(11)
    return


# li1,li2为参数列表，list1为巡检结果，list2为巡检结果说明
# list1
def clusterDocument(document, list1, list2):
    h1 = document.add_heading('2.集群巡检')
    count = 0
    text = ''
    for i in list2:
      if i:
          count += 1
    p1 = document.add_paragraph()
    run1 = p1.add_run("巡检小结：")
    run1.font.name = u'宋体'
    run1.font.size = Pt(11)
    text = "对集群虚拟化进行巡检，巡检异常项数：" + (str)(count) + "；" + "正常项数：" + (str)(len(list2) - count)
    p2 = document.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.3)
    run2 = p2.add_run(text)
    run2.font.name = u'宋体'
    run2.font.size = Pt(11)

    t1 = createTable(document, 7, 4)
    # 初始化表格
    t1.cell(0, 0).text = "检查内容"
    t1.cell(0, 1).text = "检查方法"
    t1.cell(0, 2).text = "检查结果"
    t1.cell(0, 3).text = "说明"
    t1.cell(1, 0).text = "集群高可靠性（HA）功能：查看集群的高可靠性（HA）功能是否正常开启"
    t1.cell(1, 1).text = "在<云资源>/<主机池>/<集群>的“高可靠性”页面检查是否选择了“启用HA”"
    t1.cell(2, 0).text = "集群动态资源调度（DRS）功能：查看集群的动态" \
                         "资源调度（DRS）功能是否正常开启"
    t1.cell(2, 1).text = "在<云资源>/<主机池>/<集群>的“动态资源调度”" \
                         "页面检查是否选择了“开启动态资源调度”"
    t1.cell(3, 0).text = "集群下虚拟交换机分配：查" \
                         "看集群下虚拟交换机的分配情况。"
    t1.cell(3, 1).text = "在<云资源>/<主机池>/<主机>的“虚拟交" \
                         "换机”页面检查集群下的所有主机是否都" \
                         "有相同名称的虚拟交换机"
    t1.cell(4, 0).text = "集群下共享存储分配：" \
                         "查看集群下共享存储的分配情况"
    t1.cell(4, 1).text = "在<云资源>/<主机池>/<集群>的“存储”" \
                         "页面检查集群下的主机是否都分配了相同的共享存储"
    t1.cell(5, 0).text = "集群下共享存储使用率：查看集群下共享存" \
                         "储的实际使用情况，实际使用率超过70%标记为不" \
                         "正常。实际使用率超过90%，标记为平台重大风险项。"
    t1.cell(5, 1).text = "在<云资源>/<主机池>/<集群>的“存储”页面检" \
                         "查集群下的共享存储可用容量"
    t1.cell(6, 0).text = "集群高可靠性生效最小节点数：查看集群中正常运行的主机数量不少于“HA生效最小节点数”"
    t1.cell(6, 1).text = "在<云资源>/<主机池>/<集群>的“高可靠性”页面检查“HA生效最小节点数”和集群内正常运行的主机数量"

    t1.columns[2].width = Mm(20)
    # 参数赋值
    for i in range(6):
        if not list2[i]:
            t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
        else:
            run = t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
            run.font.color.rgb = RGBColor(255, 0, 0)
            t1.cell(i + 1, 3).paragraphs[0].add_run(list2[i])
    return


def hostDocument(document, list1, list2):
    h1 = document.add_heading('3.主机巡检')
    count = 0
    text = ''
    for i in list2:
      if i:
          count += 1
    p1 = document.add_paragraph()
    run1 = p1.add_run("巡检小结：")
    run1.font.name = u'宋体'
    run1.font.size = Pt(11)
    text = "对主机CVK进行巡检，巡检异常项数：" + (str)(count) + "；" + "正常项数：" + (str)(len(list2) - count)
    p2 = document.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.3)
    run2 = p2.add_run(text)
    run2.font.name = u'宋体'
    run2.font.size = Pt(11)
    t1 = createTable(document, 8, 4)
    # 初始化表格
    t1.cell(0, 0).text = "检查内容"
    t1.cell(0, 1).text = "检查方法"
    t1.cell(0, 2).text = "检查结果"
    t1.cell(0, 3).text = "说明"
    t1.cell(1, 0).text = "*主机状态：\n*查看所有主机的运行状态。"
    t1.cell(1, 1).text = "在<云资源>的“主机”页面检查所有主" \
                         "机的运行状态是否显示“正常”"
    t1.cell(2, 0).text = "主机CPU占用率：查看所有主机CPU占用率，不超过80%"
    t1.cell(2, 1).text = "在<云资源>的“主机”页面检查所有主机的CPU占用率是否正常。"
    t1.cell(3, 0).text = "主机内存占用率：查看所有主机内存占用率，不超过80%。"
    t1.cell(3, 1).text = "在<云资源>的“主机”页面检查所有主机的内存占用率是否正常。"
    t1.cell(4, 0).text = "主机的磁盘和分区占用率：查看主机的磁盘和分区占用率，各个分区的占用率不超过80%。"
    t1.cell(4, 1).text = "在<云资源>/<主机池>/<集群>/<主机>的“性能监控”页面，查看“磁盘利用率”和“分区利用率”"
    t1.cell(5, 0).text = "主机的存储池：查看主机的存储池资源是否正常。\n*状态：活动"
    t1.cell(5, 1).text = "在<云资源>/<主机池>/<集群>/<主机>的“存储”页面，查看状态是否为“活动”，是否有足够的存储资源"
    t1.cell(6, 0).text = "主机的虚拟交换机：查看主机的虚拟交换机池资源是否正常。\n*状态：活动"
    t1.cell(6, 1).text = "在<云资源>/<主机池>/<集群>/<主机>的“虚拟交换机”页面，查看状态是否为“活动”，并且仅配置一个网关"
    t1.cell(7, 0).text = "主机的物理网卡：查看主机的物理网卡是否正常。" \
                         "\n*状态：活动\n*速率：与物理网卡实际速率保持一致" \
                         "\n*工作模式：full"
    t1.cell(7, 1).text = "在<云资源>/<主机池>/<集群>/<主机>的“物理网卡”" \
                         "页面，查看“状态”、“速率”以及“工作模式”是否正常。"

    # t1.cell(8, 0).text = "主机的FC HBA卡状态（可选）：查看主机的FC HBA卡是否" \
    #                      "正常。\n*状态：活动\n*速率：与物理FC HBA卡实际速率保持一致"
    # t1.cell(8, 1).text = "在<云资源>/<主机池>/<集群>/<主机>的“FC HBA”" \
    #                      "页面，查看“状态”和“速率”是否正常。"

    # 参数赋值
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
    for i in range(7):
        if not list2[i]:
            t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
        else:
            run = t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
            run.font.color.rgb = RGBColor(255, 0, 0)
            t1.cell(i + 1, 3).paragraphs[0].add_run(list2[i])
    return


def vmDocument(document, list1, list2):
    h1 = document.add_heading('4.虚拟机巡检')
    count = 0
    text = ''
    for i in list2:
        if i:
            count += 1
    p1 = document.add_paragraph()
    run1 = p1.add_run("巡检小结：")
    run1.font.name = u'宋体'
    run1.font.size = Pt(11)
    text = "对主机虚拟机进行巡检，巡检异常项数：" + (str)(count) + "；" + "正常项数：" + (str)(len(list2) - count)
    p2 = document.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.3)
    run2 = p2.add_run(text)
    run2.font.name = u'宋体'
    run2.font.size = Pt(11)
    t1 = createTable(document, 8, 4)
    # 初始化表格
    t1.cell(0, 0).text = "检查内容"
    t1.cell(0, 1).text = "检查方法"
    t1.cell(0, 2).text = "检查结果"
    t1.cell(0, 3).text = "说明"
    t1.cell(1, 0).text = "*虚拟机状态：\n*查看所有虚拟机的运行状态"
    t1.cell(1, 1).text = "在<云资源>的“虚拟机”页面检查所有虚拟机的运行状态。"
    t1.cell(2, 0).text = "虚拟机CPU占用率：查看所有虚拟机CPU占用率，不超过80%"
    t1.cell(2, 1).text = "在<云资源>的“虚拟机”页面检查所有主机的CPU占用率是否正常。"
    t1.cell(3, 0).text = "虚拟机内存占用率：查看所有虚拟机内存占用率，不超过80%。"
    t1.cell(3, 1).text = "在<云资源>的“虚拟机”页面检查所有虚拟机的内存占用率是否正常。"
    t1.cell(4, 0).text = "虚拟机的CAS Tools：查看虚拟机的CAS Tools工具是否正常运行。"
    t1.cell(4, 1).text = "在<云资源>/<主机池>/<集群>/<主机>/<虚拟机>的“概要”页面，查看“CAS Tools”是否为运行状态"
    t1.cell(5, 0).text = "虚拟机的磁盘和分区占用率：查看虚拟机的磁盘和分区占用率，各个分区的占用率不超过80%。"
    t1.cell(5, 1).text = "在<云资源>/<主机池>/<集群>/<主机>/<虚拟机>的“性能监控”页面，查看“磁盘利用率”和“分区利用率”"
    t1.cell(6, 0).text = "虚拟机的磁盘类型（大云可选）：查看虚拟机的磁盘信息。\n" \
                         "*设备对象：virtio磁盘 XXX\n" \
                         "*源路径：共享存储路径\n" \
                         "*缓存方式：建议使用“directsync”\n" \
                         "*存储格式：建议使用“智能”"
    t1.cell(6, 1).text = "在<云资源>/<主机池>/<集群>/<主机>/<虚拟机>的“修改虚拟机”对话框，查看“总线类型”和“存储卷路径”等"
    t1.cell(7, 0).text = "拟机的网卡（大云可选）：" \
                         "查看虚拟机的网卡信息。\n" \
                         "*设备型号：virtio网卡\n" \
                         "*内核加速：勾选"
    t1.cell(7, 1).text = "在<云资源>/<主机池>/<集群>/<主机>/<虚拟机>的“修改虚拟机”对话框，查看网卡类型。"
    # 参数赋值
    for i in range(7):
        if not list2[i]:
            t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
        else:
            run = t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
            run.font.color.rgb = RGBColor(255, 0, 0)
            t1.cell(i + 1, 3).paragraphs[0].add_run(list2[i])
    return


def systemHaDocument(document, list1, list2):
    h1 = document.add_heading('5.系统可靠性巡检')
    count = 0
    text = ''
    for i in list2:
        if i:
            count += 1
    p1 = document.add_paragraph()
    run1 = p1.add_run("巡检小结：")
    run1.font.name = u'宋体'
    run1.font.size = Pt(11)
    text = "对系统可靠性进行巡检，巡检异常项数：" + (str)(count) + "；" + "正常项数：" + (str)(len(list2) - count)
    p2 = document.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.3)
    run2 = p2.add_run(text)
    run2.font.name = u'宋体'
    run2.font.size = Pt(11)
    t1 = createTable(document, 5, 4)
    t1.style.font.name = '微软雅黑'
    t1.style.font.size = Pt(9)
    # 初始化表格
    t1.cell(0, 0).text = "检查内容"
    t1.cell(0, 1).text = "检查方法"
    t1.cell(0, 2).text = "检查结果"
    t1.cell(0, 3).text = "说明"
    t1.cell(1, 0).text = "链路冗余：查看系统的链路冗余情况。"
    t1.cell(1, 1).text = "在<云资源>/<主机池>/<集群>/<主机>/<虚拟交换机>页面，检查各个虚拟交换机是否进行了链路冗余（动态或者静态聚合）"
    t1.cell(2, 0).text = "CVM配置备份：查看CVM配置的备份情况。用户CVM主机故障时的系统配置恢复。"
    t1.cell(2, 1).text = "在<系统管理>/<安全管理>的“CVM备份配置”页面，确认已启用了定时备份功能，推荐备份到远端目录。"
    t1.cell(3, 0).text = "CVM双机热备状态检查：检查CAS的CVM双机热备状态是否正常。"
    t1.cell(3, 1).text = "在CVM双机热备环境中，随意登录CVM主机后台执行“crm status”检查双机热备状态。"
    t1.cell(4, 0).text = "虚拟机的备份：检查重要虚拟机是否已经开启备份功能。"
    t1.cell(4, 1).text = "检查运行客户重要业务的虚拟机是否开启了定时备份功能。"
    # 参数赋值
    for i in range(4):
        if not list2[i]:
            t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
        else:
            run = t1.cell(i + 1, 2).paragraphs[0].add_run(list1[i])
            run.font.color.rgb = RGBColor(255, 0, 0)
            t1.cell(i + 1, 3).paragraphs[0].add_run(list2[i])
    return

# cvm平台信息巡检
def cvmCheck(document, casInfo):
    list1 = []
    list1.append(casInfo['productVersion'])
    list1.append(casInfo['deviceDmide'])
    list1.append(casInfo['casVersion'])
    list1.append(casInfo['installType'])
    list1.append(casInfo['ovsVersion'])
    list1.append(casInfo['licenseInfo'])
    casBasicDocument(document, list1)
    del list1
    return


###################
# 集群巡检        #
##################
def clusterCheck(document, casInfo):
    list1 = []
    list2 = ['' for n in range(7)]

    # 集群是否开启HA和DRS
    tempHa = ''
    tempLB = ''
    for i in casInfo['clusterInfo']:
        if i['enableHA'] == '0':
            list2[0] += "集群" + i['name'] + " HA未开启\n"
        if i['enableLB'] == '0':
            list2[1] = "集群" + i['name'] + " DRS未开启\n"

    # 集群下主机虚拟交换机部署是否合规
    dict1 = dict()
    for i in casInfo['clusterInfo']:
        dict1[i['name']] = list()
        for j in i['cvkInfo']:
            for k in j['vswitch']:
                if not k['name'] in dict1[i['name']]:
                    dict1[i['name']].append(k['name'])
    for i in casInfo['clusterInfo']:
        if len(dict1[i['name']]) != 3:
            list2[2] += "集群" + i['name'] + "下交换机的部署不合规\n"

    # cvk共享存储池部署是否一致
    dict1 = {}  # 存储集群下的所有共享存储池
    dict2 = {}  # 存储主机下的共享存储池
    for i in casInfo['clusterInfo']:
        dict1[i['name']] = set()
        for j in i['cvkInfo']:
            dict2[j['name']] = set()
            for k in j['sharePool']:
                dict1[i['name']].add(k['name'])
                dict2[j['name']].add(k['name'])

        for m in i['cvkInfo']:
            if dict1[i['name']] != dict2[m['name']]:
                list2[3] += "集群" + i['name'] + "下主机" + m['name'] + "共享存储池与集群不一致"
    del dict1, dict2

    # 共享存储利用率：
    for i in casInfo['clusterInfo']:
        li1 = list()
        for j in i['cvkInfo']:
            for k in j['sharePool']:
                if not k in li1:
                    li1.append(k)
        for h in li1:
            if h['rate'] > 0.8:
                list2[4] = "集群" + i['name'] + "下共享存储池" + h['name'] + "利用率超过80%达到" + str(h['rate'])
        del li1

    # 集群最小主机节点
    for i in casInfo['clusterInfo']:
        if i['enableHA'] == '0':
            list2[5] = "集群未开启高可靠"
        else:
            if (int)(i['HaMinHost']) > len(i['cvkInfo']):
                list2[5] = "Ha最小节点数小正常运行主机数"

    for i in list2:
        if not i:
            list1.append("正常")
        else:
            list1.append("异常")

    clusterDocument(document, list1, list2)
    del list1, list2
    return


#######################################
# 主机巡检                            #
#                                     #
########################################
def cvkCheck(document, casInfo):
    list1 = []
    list2 = ['' for n in range(7)]
    for i in casInfo['clusterInfo']:
        dict1 = {}
        for j in i['cvkInfo']:
            dict1[j['name']] = ''
            # 主机状态检测
            if j['status'] != '1':
                if not list2[0]:
                    list2[0] += "状态异常主机如下" + j['name'] + '\n'
                else:
                    list2[0] += j['name'] + '\n'

            # 主机cpu利用率
            if j['cpuRate'] > 80:
                if not list2[1]:
                    list2[1] += "cpu利用率超过80%主机如下：" + j['name'] + '\n'
                else:
                    list2[1] += j['name'] + '\n'

            # 主机内存利用率
            if j['memRate'] > 80:
                if not list2[2]:
                    list2[2] += "内存利用率超过80%主机如下：" + j['name'] + '\n'
                else:
                    list2[2] += j['name'] + '\n'

            # 主机磁盘利用率
            for k in j['diskRate']:
                if (float)(k['usage']) > 80:
                    if not dict1[j['name']]:
                        dict1[j['name']] += "\n主机" + j["name"] + "磁盘利用率查过80%的磁盘如下：" + k["name"]
                    else:
                        dict1[j['name']] += "、" + k["name"]
            for h in dict1.values():
                list2[3] += h

            # 主机存储池状态：
            for m in j['storagePool']:
                if m['status'] != '1':
                    if not list2[4]:
                        list2[4] = "\n主机" + j['name'] + "状态异常磁盘如下：" + m['name']
                    else:
                        list2[4] += m['name']

            # 虚拟交换机状态
            for k in j['vswitch']:
                if k['status'] != '1':
                    if not list2[5]:
                        list2[5] = "\n主机" + j['name'] + "状态异常虚拟交换机如下：" + k['name']
                    else:
                        list2[5] += k['name']

            # 网卡状态
            for k in j['network']:
                if k['status'] != 'yes' or (int)(k['speed']) < 1000 or k['duplex'] != 'Full':
                    if not list2[6]:
                        list2[6] = "\n主机" + j['name'] + "状态异常网卡如下：" + k['name']
                    else:
                        list2[6] += k['name']
        del dict1
    # 主机巡检结果写入docx
    for i in list2:
        if not i:
            list1.append("正常")
        else:
            list1.append("异常")
    hostDocument(document, list1, list2)
    del list1, list2
    return


#######################################
# 虚拟机巡检                            #
#                                      #
########################################

def vmCheck(document, casInfo):
    list1 = []
    list2 = ['' for n in range(7)]
    for i in casInfo['clusterInfo']:
        for j in i['cvkInfo']:
            dict1 = {}
            dict2 = {}
            for k in j['vmInfo']:
                # print(k)
                # 虚拟机状态
                if k['status'] != 'running':
                    if not list2[0]:
                        list2[0] = "状态异常虚拟机如下：" + k['name']
                    else:
                        list2[0] += '、' + k['name']
                else:
                    # 虚拟机cpu利用率是否超过80%
                    if k['cpuReate'] > 80:
                        if not list2[1]:
                            list2[1] = "cpu利用率超过80%虚拟机如下：" + k['name']
                        else:
                            list2[1] += '、' + k['name']

                    # 虚拟机内存利用率是否超过80%
                    if k['memRate'] > 80:
                        if not list2[2]:
                            list2[2] = "内存利用率超过80%虚拟机如下：" + k['name']
                        else:
                            list2[2] += '、' + k['name']

                    # 虚拟机castool状态异常
                    if k['castoolsStatus'] != '1':
                        if not list2[3]:
                            list2[3] = "castool状态虚拟机如下：" + k['name']
                        else:
                            list2[3] += '、' + k['name']

                    # 虚拟机磁盘分区巡检
                    tmp = ''
                    for m in k['diskRate']:
                        if m['usage'] > 80:
                            if not tmp:
                                tmp = "\n虚拟机" + k['name'] + '磁盘利用率超过80%的磁盘如下：' + m['name']
                            else:
                                tmp += '、' + m['name']
                    list2[4] += tmp
                    del tmp

                    # 虚拟机磁盘巡检
                    dict1[k['name']] = ''
                    for n in k['vmdisk']:
                        tmp = n['path'].split('/')
                        path = '/' + tmp[1] + '/' + tmp[2]
                        bool1 = False
                        for m in j['sharePool']:
                            if path == m['path']:
                                bool1 = True
                        if not bool1:
                            if not dict1[k['name']]:
                                dict1[k['name']] = "\n虚拟机" + k['name'] + "磁盘" + n['name'] + '使用了非共享存储池'
                            else:
                                dict1[k['name']] += "磁盘" + n['name'] + '使用了非共享存储池'
                        if n['format'] != 'qcow2':
                            if not dict1[k['name']]:
                                dict1[k['name']] = "\n虚拟机" + k['name'] + "磁盘" + n['name'] + '格式错误'
                            else:
                                dict1[k['name']] += "磁盘" + n['name'] + '格式错误'
                        if n['cacheType'] != 'directsync':
                            if not dict1[k['name']]:
                                dict1[k['name']] = "\n虚拟机" + k['name'] + "磁盘" + n['name'] + '缓存方式错误'
                            else:
                                dict1[k['name']] += "磁盘" + n['name'] + '缓存方式错误'
                    list2[5] += dict1[k['name']]

                    # 虚拟机网卡巡检
                    dict2[k['name']] = ''
                    for m in k['vmNetwork']:
                        if m['mode'] != 'virtio':
                            if not dict2[k['name']]:
                                dict2[k['name']] = '\n虚拟机' + k['name'] + '网卡' + m['name'] + '模式错误'
                            else:
                                dict2[k['name']] = '网卡' + m['name'] + '模式错误'
                        if m['KernelAccelerated'] != '1':
                            if not dict2[k['name']]:
                                dict2[k['name']] = '\n虚拟机' + k['name'] + '网卡' + m['name'] + '未开启内核加速'
                            else:
                                dict2[k['name']] = '网卡' + m['name'] + '未开启内核加速'
                    list2[6] += dict2[k['name']]
            del dict1, dict2

    for i in list2:
        if not i:
            list1.append("正常")
        else:
            list1.append("异常")
    vmDocument(document, list1, list2)
    del list1, list2
    return


####################
# cvm可靠性巡检
####################
def cvmHaCheck(document, casInfo):
    list1 = []
    list2 = ['' for n in range(4)]

    # 虚拟交换机的是否配置冗余链路

    # cvm是否开启备份策略
    if not casInfo['BackupEnable']:
        list2[1] = 'cvm未开启备份策略'

    # cvm是否开启HA高可靠
    if not casInfo['HA']:
        list2[2] = '未开启HA高可靠'

    # 检查虚拟机是否配置高可靠
    if casInfo['vmBackPolicy'] == 'NONE':
        list2[3] = '未配置虚拟机备份'
    else:
        for i in casInfo['vmBackPolicy']:
            if i['state'] != '1':
                if not list2[3]:
                    list2[3] = '状态异常备份策略如下：' + i['name']
                else:
                    list2[3] += '、' + i['name']

    for i in list2:
        if not i:
            list1.append("正常")
        else:
            list1.append("异常")
    systemHaDocument(document, list1, list2)
    return



# if __name__ == '__main__':
#     cas = casCheck('192.168.2.5', 'admin', 'admin', 'root', 'h3c.com!')
#     print(cas.casInfo)
#     for i in cas.casInfo['clusterInfo']:
#         for j in i['cvkInfo']:
#             print(j)
    # document = openDocument()
    # cvmCheck(document, cas)
    # clusterCheck(document, cas)
    # cvkCheck(document, cas)
    # vmCheck(document, cas)
    # cvmHaChech(document, cas)
    # document.save('test1.docx')
